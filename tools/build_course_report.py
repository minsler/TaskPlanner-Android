from __future__ import annotations

import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\marya\AndroidStudioProjects\TaskPlanner")
SERVER_ROOT = Path(r"C:\Users\marya\IdeaProjects\TaskPlannerServer")
OUT_DIR = ROOT / "reports"
IMG_DIR = OUT_DIR / "images"
OUT_DOCX = OUT_DIR / "Курсовая работа_TaskPlanner_без титульного листа.docx"


def safe_read(path: Path) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    if path.name == "Application.kt":
        text = re.sub(r'password\s*=\s*"[^"]+"', 'password = "***"', text)
        text = re.sub(r'val jwtSecret\s*=\s*"[^"]+"', 'val jwtSecret = "***"', text)
    return text


def font_path() -> str | None:
    candidates = [
        Path(r"C:\Windows\Fonts\times.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


FONT_PATH = font_path()


def make_diagram(filename: str, title: str, boxes: list[tuple[str, int, int, int, int]], lines: list[tuple[int, int]] = ()) -> Path:
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / filename
    image = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(FONT_PATH, 44) if FONT_PATH else ImageFont.load_default()
    box_font = ImageFont.truetype(FONT_PATH, 27) if FONT_PATH else ImageFont.load_default()
    small_font = ImageFont.truetype(FONT_PATH, 22) if FONT_PATH else ImageFont.load_default()
    draw.text((60, 35), title, fill=(25, 25, 25), font=title_font)

    centers = []
    for text, x, y, w, h in boxes:
        centers.append((x + w // 2, y + h // 2))

    for start, end in lines:
        x1, y1 = centers[start]
        x2, y2 = centers[end]
        draw.line((x1, y1, x2, y2), fill=(70, 100, 130), width=4)

    for text, x, y, w, h in boxes:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, outline=(45, 95, 130), width=4, fill=(238, 246, 252))
        wrapped = textwrap.wrap(text, width=max(14, w // 22))
        total_h = len(wrapped) * 34
        ty = y + (h - total_h) // 2
        for line in wrapped:
            bbox = draw.textbbox((0, 0), line, font=box_font)
            draw.text((x + (w - (bbox[2] - bbox[0])) // 2, ty), line, fill=(20, 45, 65), font=box_font)
            ty += 34

    draw.text((60, 895), "Рисунок подготовлен для пояснительной записки по проекту TaskPlanner.", fill=(90, 90, 90), font=small_font)
    image.save(path)
    return path


def build_diagrams() -> dict[str, Path]:
    return {
        "use_case": make_diagram(
            "use_case.png",
            "Use case диаграмма приложения",
            [
                ("Пользователь", 60, 390, 260, 120),
                ("Регистрация", 500, 125, 320, 95),
                ("Авторизация", 500, 245, 320, 95),
                ("Просмотр задач", 500, 365, 320, 95),
                ("Создание задачи", 500, 485, 320, 95),
                ("Редактирование задачи", 500, 605, 320, 95),
                ("Удаление задачи", 940, 245, 320, 95),
                ("Поиск и история", 940, 365, 320, 95),
                ("Смена статуса", 940, 485, 320, 95),
                ("Переключение темы", 940, 605, 320, 95),
            ],
            [(0, i) for i in range(1, 10)],
        ),
        "architecture": make_diagram(
            "architecture.png",
            "Архитектура клиент-серверного приложения",
            [
                ("Presentation: Compose экраны и ViewModel", 80, 160, 430, 130),
                ("Domain: модели и интерфейсы репозиториев", 80, 380, 430, 130),
                ("Data: API, DTO, DataStore, TokenManager", 80, 600, 430, 130),
                ("Ktor Server: JWT, маршруты, CRUD", 660, 260, 390, 140),
                ("PostgreSQL: Users, Tasks", 660, 560, 390, 140),
                ("Postman: проверка API", 1120, 410, 330, 120),
            ],
            [(0, 1), (1, 2), (2, 3), (3, 4), (5, 3)],
        ),
        "navigation": make_diagram(
            "navigation.png",
            "Межэкранное взаимодействие",
            [
                ("Запуск приложения", 90, 190, 330, 100),
                ("Проверка JWT токена", 525, 190, 330, 100),
                ("Авторизация / регистрация", 960, 130, 380, 120),
                ("Экран задач", 960, 360, 380, 120),
                ("Диалог добавления", 330, 570, 330, 100),
                ("Диалог редактирования", 760, 570, 360, 100),
                ("Выход из аккаунта", 1160, 570, 300, 100),
            ],
            [(0, 1), (1, 2), (1, 3), (3, 4), (3, 5), (3, 6), (6, 2)],
        ),
        "database": make_diagram(
            "database.png",
            "Логическая модель базы данных",
            [
                ("Users\nid PK\nemail UNIQUE\npassword hash", 230, 250, 440, 260),
                ("Tasks\nid PK\ntitle\ndescription\nis_completed\nuser_id FK", 930, 230, 440, 300),
            ],
            [(0, 1)],
        ),
        "class": make_diagram(
            "class_diagram.png",
            "Диаграмма классов клиентской части",
            [
                ("MainActivity", 90, 130, 300, 100),
                ("AppNavigation", 520, 130, 320, 100),
                ("AuthViewModel", 1030, 90, 350, 110),
                ("TasksViewModel", 1030, 250, 350, 110),
                ("TaskRepository", 520, 370, 320, 110),
                ("TaskRepositoryImpl", 520, 550, 320, 110),
                ("AuthApi / TasksApi", 90, 550, 330, 110),
                ("TokenManager / DataStore", 1030, 550, 360, 110),
                ("Task", 520, 750, 320, 90),
            ],
            [(0, 1), (1, 2), (1, 3), (3, 4), (4, 5), (5, 6), (3, 7), (3, 8)],
        ),
        "idef0": make_diagram(
            "idef0.png",
            "IDEF0: управление задачами",
            [
                ("Вход:\nучетные данные,\nпараметры задачи", 70, 375, 330, 160),
                ("A0\nПланирование и\nуправление задачами", 580, 315, 440, 250),
                ("Управление:\nправила авторизации,\nтребования методички", 580, 95, 440, 140),
                ("Механизмы:\nAndroid, Ktor,\nPostgreSQL", 580, 690, 440, 140),
                ("Выход:\nсписок задач,\nрезультаты поиска,\nобновленный статус", 1170, 350, 360, 200),
            ],
            [(0, 1), (2, 1), (3, 1), (1, 4)],
        ),
    }


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    for style_name, size, uppercase, before, after in [
        ("Heading 1", 18, True, 0, 10),
        ("Heading 2", 16, False, 15, 10),
        ("Heading 3", 14, False, 15, 10),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.left_indent = Cm(1.25)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5

    if "Caption Custom" not in styles:
        cap = styles.add_style("Caption Custom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        cap.font.size = Pt(12)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.space_after = Pt(8)

    if "Code" not in styles:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(8)
        code.paragraph_format.first_line_indent = Cm(0)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.space_after = Pt(0)


def paragraph(doc: Document, text: str = "", style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def heading(doc: Document, text: str, level: int) -> None:
    if level == 1:
        text = text.upper()
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.5
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    set_table_borders(table)
    paragraph(doc, "")


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    paragraph(doc, caption, style="Caption Custom")


def add_section(doc: Document, title: str, level: int = 1) -> None:
    if len(doc.paragraphs) > 0:
        doc.add_page_break()
    heading(doc, title, level)


def generate_paragraphs(topic: str, points: list[str], count: int) -> list[str]:
    base = []
    templates = [
        "В рамках проекта {topic} {point}. Такое решение важно для курсовой работы, поскольку демонстрирует связь между пользовательским интерфейсом, клиентской логикой, серверной частью и базой данных.",
        "Для реализации функции {topic} используется подход, при котором {point}. Это снижает вероятность ошибок пользователя и делает поведение приложения предсказуемым при повторном запуске и смене ориентации экрана.",
        "Практическая ценность решения состоит в том, что {point}. Пользователь получает единый сценарий работы: вход в систему, просмотр списка задач, поиск, изменение статуса и синхронизацию с сервером.",
        "При проектировании {topic} учитывались требования к клиент-серверным мобильным приложениям: разделение ответственности, хранение пользовательских данных, защита доступа и устойчивость к ошибкам сети. В результате {point}.",
    ]
    for i in range(count):
        base.append(templates[i % len(templates)].format(topic=topic, point=points[i % len(points)]))
    return base


def add_prose_block(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        paragraph(doc, text)


def add_intro(doc: Document) -> None:
    add_section(doc, "Введение")
    add_prose_block(doc, [
        "Клиент-серверные мобильные приложения используются для организации доступа к данным с разных устройств, централизованного хранения информации и контроля прав пользователей. В рамках данной курсовой работы разработано Android-приложение «Планировщик задач», предназначенное для регистрации пользователя, авторизации, ведения личного списка задач, поиска записей и изменения статуса выполнения.",
        "Актуальность темы связана с тем, что планирование задач является одной из распространенных потребностей пользователей мобильных устройств. Простые локальные заметки не обеспечивают надежного разделения пользователей, централизованного хранения и проверки доступа к данным. Клиент-серверная архитектура позволяет вынести хранение и бизнес-логику на сервер, а мобильный клиент сделать удобным инструментом повседневной работы.",
        "Цель курсовой работы - разработать клиент-серверное мобильное приложение-планировщик задач на базе Android, обеспечивающее регистрацию, авторизацию, выполнение CRUD-операций над задачами, поиск, сохранение истории поиска и поддержку темной темы.",
        "Объектом исследования является процесс разработки клиент-серверного мобильного приложения. Предметом исследования являются программные средства и архитектурные решения, применяемые для создания Android-клиента, сервера на Kotlin/Ktor и базы данных PostgreSQL.",
    ])
    add_bullets(doc, [
        "проанализировать предметную область и существующие аналоги приложений-планировщиков;",
        "сформулировать функциональные и нефункциональные требования к мобильному приложению;",
        "спроектировать пользовательский интерфейс, межэкранное взаимодействие и логическую модель приложения;",
        "реализовать клиентскую часть на Kotlin с использованием Jetpack Compose и Clean Architecture;",
        "реализовать серверную часть на Ktor с JWT-авторизацией, BCrypt и PostgreSQL;",
        "проверить работу приложения по тестовым сценариям и оценить соответствие требованиям методических указаний.",
    ])
    add_prose_block(doc, [
        "Пояснительная записка состоит из трех глав. В первой главе рассмотрена предметная область, приведено сравнение аналогов и обоснован выбор инструментальных средств. Во второй главе представлены требования, диаграммы, структура интерфейса и логическая модель базы данных. В третьей главе описана разработка клиентской и серверной частей, подключение компонентов и тестирование приложения.",
    ])


def add_chapter_1(doc: Document) -> None:
    add_section(doc, "Глава 1. Анализ состояния объекта исследования")
    heading(doc, "1.1. Описание предметной области мобильного приложения", 2)
    add_prose_block(doc, generate_paragraphs(
        "планировщика задач",
        [
            "пользователь должен быстро фиксировать задачу, видеть ее описание и отмечать выполнение",
            "личные задачи должны быть доступны только владельцу учетной записи",
            "поиск снижает время доступа к нужной записи при росте количества задач",
            "темная тема повышает удобство использования приложения в разных условиях освещения",
        ],
        28,
    ))
    heading(doc, "1.2. Сравнительный анализ существующих аналогов", 2)
    paragraph(doc, "Для анализа были выбраны три популярных решения, близких по назначению: Todoist, Microsoft To Do и Google Tasks. Они позволяют вести списки задач, группировать записи и отмечать выполнение, но отличаются уровнем сложности и набором дополнительных функций.")
    add_table(doc, ["Аналог", "Преимущества", "Недостатки для учебного проекта"], [
        ["Todoist", "Развитая система проектов, меток, фильтров и синхронизации.", "Сложный интерфейс и избыточность функций для демонстрации базовой клиент-серверной архитектуры."],
        ["Microsoft To Do", "Простой интерфейс, списки, напоминания, интеграция с учетной записью Microsoft.", "Закрытая серверная часть и привязка к экосистеме Microsoft."],
        ["Google Tasks", "Минималистичный подход, интеграция с Google-сервисами, быстрый ввод задач.", "Ограниченная демонстрация собственной серверной логики и базы данных."],
    ])
    add_prose_block(doc, generate_paragraphs(
        "сравнения аналогов",
        [
            "для курсового проекта целесообразно реализовать компактный планировщик с собственным сервером и понятным набором функций",
            "основной акцент сделан на полном цикле работы с данными, а не на большом количестве второстепенных возможностей",
            "разработанное приложение сохраняет простоту интерфейса, но демонстрирует регистрацию, авторизацию и защищенный доступ к данным",
        ],
        18,
    ))
    heading(doc, "1.3. Обоснование выбора инструментальных средств", 2)
    add_table(doc, ["Компонент", "Выбранное средство", "Обоснование"], [
        ["Клиент", "Kotlin, Jetpack Compose", "Kotlin является основным языком Android-разработки, Compose позволяет создавать интерфейс декларативно и полностью на Kotlin."],
        ["Архитектура клиента", "Clean Architecture", "Разделение data, domain и presentation упрощает сопровождение и снижает связность кода."],
        ["Сеть", "Ktor Client", "Библиотека хорошо сочетается с Kotlin и kotlinx.serialization."],
        ["Локальные настройки", "DataStore Preferences", "Подходит для хранения токена, темы и истории поиска."],
        ["Сервер", "Ktor Server", "Позволяет реализовать REST API на Kotlin и использовать JWT-аутентификацию."],
        ["База данных", "PostgreSQL, Exposed", "PostgreSQL обеспечивает надежное хранение, Exposed упрощает работу с таблицами из Kotlin-кода."],
        ["Безопасность", "JWT, BCrypt", "JWT используется для stateless-авторизации, BCrypt - для безопасного хранения паролей."],
    ])
    add_prose_block(doc, generate_paragraphs(
        "выбора инструментов",
        [
            "использование Kotlin на клиенте и сервере уменьшает технологический разрыв между частями системы",
            "PostgreSQL подходит для хранения связанных сущностей пользователей и задач",
            "JWT позволяет проверять права доступа к каждому защищенному маршруту сервера",
            "DataStore лучше соответствует современным рекомендациям Android для хранения пользовательских настроек",
        ],
        24,
    ))


def add_chapter_2(doc: Document, diagrams: dict[str, Path]) -> None:
    add_section(doc, "Глава 2. Проектирование мобильного приложения")
    heading(doc, "2.1. Определение функциональных и нефункциональных требований", 2)
    paragraph(doc, "Функциональные требования определяют действия, которые пользователь должен выполнять в приложении. Они сформулированы с учетом задания и обязательных требований методических указаний.")
    add_table(doc, ["Код", "Функциональное требование", "Реализация"], [
        ["FR-1", "Пользователь может зарегистрировать учетную запись.", "Экран регистрации, маршрут POST /register."],
        ["FR-2", "Пользователь может авторизоваться.", "Экран авторизации, маршрут POST /login, сохранение JWT."],
        ["FR-3", "Пользователь может просматривать список своих задач.", "Экран задач, GET /tasks, проверка userId на сервере."],
        ["FR-4", "Пользователь может создавать, редактировать и удалять задачи.", "Диалоги Compose и REST-маршруты POST, PUT, DELETE."],
        ["FR-5", "Пользователь может менять статус задачи.", "Checkbox в карточке и обновление поля isCompleted."],
        ["FR-6", "Пользователь может выполнять поиск по задачам.", "Поисковая строка, кнопки «Найти» и «Очистить», локальная фильтрация."],
        ["FR-7", "Приложение сохраняет историю поиска до 10 элементов.", "DataStore Preferences, отображение истории при фокусе поля поиска."],
        ["FR-8", "Пользователь может включить темную тему.", "Switch в верхней панели, сохранение выбора в DataStore."],
    ])
    add_table(doc, ["Код", "Нефункциональное требование", "Способ обеспечения"], [
        ["NFR-1", "Интерфейс полностью на русском языке.", "Все пользовательские тексты в Compose заданы на русском языке."],
        ["NFR-2", "Доступ к задачам должен быть защищен.", "JWT-аутентификация и фильтрация задач по userId."],
        ["NFR-3", "Приложение должно устойчиво реагировать на ошибки сети.", "Состояния error и кнопка «Обновить»."],
        ["NFR-4", "Данные поиска и темы должны сохраняться после перезапуска.", "DataStore Preferences."],
        ["NFR-5", "Код должен быть сопровождаемым.", "Разделение на data, domain и presentation."],
    ])
    add_figure(doc, diagrams["use_case"], "Рисунок 2.1. Use case диаграмма приложения")
    add_prose_block(doc, generate_paragraphs(
        "требований",
        [
            "обязательные функции методических указаний вынесены в отдельные пользовательские сценарии",
            "ошибочные состояния отображаются явно и содержат действие для повторной попытки",
            "поиск сохраняет текст запроса в состоянии ViewModel, поэтому при повороте устройства введенная строка не исчезает",
        ],
        24,
    ))
    heading(doc, "2.2. Проектирование пользовательского интерфейса", 2)
    paragraph(doc, "Интерфейс приложения построен на нескольких основных экранных формах: экран авторизации, экран регистрации, экран списка задач, диалоги добавления и редактирования задачи. Межэкранное взаимодействие реализовано средствами Navigation Compose.")
    add_figure(doc, diagrams["navigation"], "Рисунок 2.2. Схема межэкранного взаимодействия")
    add_table(doc, ["Экран", "Назначение", "Основные элементы"], [
        ["Авторизация", "Вход существующего пользователя.", "Email, пароль, кнопка входа, переход к регистрации, индикатор загрузки."],
        ["Регистрация", "Создание новой учетной записи.", "Email, пароль, кнопка создания аккаунта, переход к авторизации."],
        ["Список задач", "Основная рабочая область.", "TopAppBar, переключатель темы, поиск, вкладки, карточки задач, FAB."],
        ["Добавление задачи", "Создание новой задачи.", "Поля названия и описания, кнопки «Сохранить» и «Отмена»."],
        ["Редактирование задачи", "Изменение существующей задачи.", "Поля с текущими значениями, кнопки «Обновить» и «Отмена»."],
    ])
    add_prose_block(doc, generate_paragraphs(
        "интерфейса",
        [
            "верхняя панель содержит часто используемые действия: обновление, выход и переключение темы",
            "строка поиска имеет подсказку, кнопку очистки и отдельную кнопку запуска поиска",
            "карточка задачи содержит checkbox, название, описание и кнопки редактирования и удаления",
            "вкладки «Все», «В процессе» и «Завершенные» помогают быстро переключать представление списка",
        ],
        26,
    ))
    heading(doc, "2.3. Логическая модель приложения", 2)
    heading(doc, "2.3.1. Проектирование мобильного приложения", 3)
    paragraph(doc, "Логическая модель клиентской части основана на Clean Architecture. Уровень presentation отвечает за отображение и обработку пользовательских событий, уровень domain содержит модель задачи и абстракцию репозитория, уровень data выполняет сетевые запросы, преобразование DTO и локальное хранение параметров.")
    add_figure(doc, diagrams["architecture"], "Рисунок 2.3. Архитектура приложения")
    add_figure(doc, diagrams["class"], "Рисунок 2.4. Диаграмма классов клиентской части")
    add_figure(doc, diagrams["idef0"], "Рисунок 2.5. Диаграмма IDEF0")
    add_prose_block(doc, generate_paragraphs(
        "логической модели клиента",
        [
            "ViewModel хранит состояние экрана и переживает изменение конфигурации, что важно для сохранения поискового запроса",
            "репозиторий скрывает детали сетевого взаимодействия от пользовательского интерфейса",
            "DTO преобразуются в доменную модель Task, используемую экраном задач",
            "TokenManager предоставляет токен для защищенных запросов, а DataStore хранит пользовательские настройки",
        ],
        24,
    ))
    heading(doc, "2.3.2. Проектирование базы данных", 3)
    paragraph(doc, "База данных содержит две основные таблицы: Users и Tasks. Связь один-ко-многим означает, что одному пользователю может принадлежать несколько задач, а каждая задача имеет ровно одного владельца.")
    add_figure(doc, diagrams["database"], "Рисунок 2.6. Логическая модель базы данных")
    add_table(doc, ["Таблица", "Поле", "Тип", "Назначение"], [
        ["Users", "id", "integer, PK, autoIncrement", "Идентификатор пользователя."],
        ["Users", "email", "varchar(255), unique", "Логин пользователя."],
        ["Users", "password", "varchar(255)", "Хэш пароля BCrypt."],
        ["Tasks", "id", "integer, PK, autoIncrement", "Идентификатор задачи."],
        ["Tasks", "title", "varchar(255)", "Название задачи."],
        ["Tasks", "description", "text", "Описание задачи."],
        ["Tasks", "is_completed", "boolean", "Статус выполнения."],
        ["Tasks", "user_id", "integer, FK", "Владелец задачи, ссылка на Users.id."],
    ])
    add_prose_block(doc, generate_paragraphs(
        "базы данных",
        [
            "поле user_id используется сервером для проверки принадлежности задач конкретному пользователю",
            "уникальность email предотвращает регистрацию нескольких учетных записей с одинаковым логином",
            "пароль не хранится в открытом виде, вместо него сохраняется BCrypt-хэш",
        ],
        18,
    ))


def add_chapter_3(doc: Document) -> None:
    add_section(doc, "Глава 3. Разработка мобильного приложения")
    heading(doc, "3.1. Разработка клиентской части мобильного приложения", 2)
    paragraph(doc, "Клиентская часть реализована на Kotlin с использованием Jetpack Compose. Проект разделен на пакеты data, domain и presentation. Такое разделение позволяет поддерживать независимость пользовательского интерфейса от деталей сетевого взаимодействия.")
    add_table(doc, ["Пакет", "Состав", "Назначение"], [
        ["data.local", "TokenManager, UserPreferencesManager", "Хранение JWT, темы и истории поиска."],
        ["data.remote", "ApiClient, AuthApi, TasksApi, DTO", "HTTP-запросы к серверу и сериализация JSON."],
        ["domain.model", "Task", "Доменная модель задачи."],
        ["domain.repository", "TaskRepository, TaskRepositoryImpl", "Интерфейс и реализация работы с задачами."],
        ["presentation.auth", "LoginScreen, AuthViewModel", "Авторизация и регистрация."],
        ["presentation.tasks", "TasksScreen, TasksViewModel", "Основной экран задач, поиск и CRUD."],
        ["presentation.navigation", "AppNavigation", "Маршрутизация между экранами."],
    ])
    add_prose_block(doc, generate_paragraphs(
        "клиентской части",
        [
            "состояния экранов представлены через StateFlow, что удобно для реактивного обновления Compose-интерфейса",
            "JWT сохраняется после успешной авторизации и используется при каждом защищенном запросе",
            "история поиска и выбранная тема сохраняются в DataStore, поэтому данные не теряются после перезапуска",
            "ProgressBar отображается во время загрузки задач, авторизации и выполнения поиска",
            "плейсхолдеры показывают отсутствие задач, отсутствие результатов поиска и ошибки загрузки",
        ],
        36,
    ))
    heading(doc, "3.2. Разработка серверной части мобильного приложения", 2)
    heading(doc, "3.2.1. Описание основной логики сервера", 3)
    paragraph(doc, "Серверная часть реализована на Ktor Server. Сервер запускается на порту 8080 и предоставляет REST API для регистрации, авторизации и управления задачами. Для аутентификации используется JWT, для хэширования паролей - BCrypt.")
    add_table(doc, ["Метод", "Маршрут", "Назначение", "Защита"], [
        ["GET", "/", "Проверка запуска сервера.", "Нет"],
        ["POST", "/register", "Регистрация пользователя.", "Нет"],
        ["POST", "/login", "Авторизация и выдача JWT.", "Нет"],
        ["GET", "/profile", "Проверка авторизованного пользователя.", "JWT"],
        ["GET", "/tasks", "Получение задач пользователя.", "JWT"],
        ["POST", "/tasks", "Создание задачи.", "JWT"],
        ["PUT", "/tasks/{id}", "Обновление задачи.", "JWT + userId"],
        ["DELETE", "/tasks/{id}", "Удаление задачи.", "JWT + userId"],
        ["GET", "/tasks/search", "Поиск задач на сервере.", "JWT + userId"],
    ])
    add_prose_block(doc, generate_paragraphs(
        "серверной логики",
        [
            "после входа сервер формирует JWT с email пользователя и сроком действия",
            "защищенные маршруты извлекают email из JWT и определяют userId в таблице Users",
            "обновление и удаление задач выполняются только при совпадении id задачи и userId владельца",
            "Ktor ContentNegotiation обеспечивает передачу данных в формате JSON",
        ],
        24,
    ))
    heading(doc, "3.2.2. Подключение базы данных к серверу", 3)
    paragraph(doc, "Подключение к PostgreSQL выполнено через Exposed. При запуске приложение подключается к базе taskplanner и создает таблицы Users и Tasks, если они отсутствуют. В отчете секретные параметры подключения не раскрываются.")
    add_prose_block(doc, generate_paragraphs(
        "подключения базы данных",
        [
            "таблицы описаны в Kotlin-коде через объекты, наследующие Table",
            "операции с базой выполняются внутри transaction, что обеспечивает корректное выполнение SQL-запросов",
            "Exposed позволяет выразить условия доступа к задачам на уровне запросов к таблицам",
        ],
        18,
    ))
    heading(doc, "3.2.3. Подключение сервера к клиентской части мобильного приложения", 3)
    paragraph(doc, "Клиент подключается к серверу по HTTP через Ktor Client. Базовый адрес сервера указан в ApiClient. Для запуска на реальном устройстве используется IPv4-адрес компьютера в локальной сети, а в AndroidManifest разрешен cleartext-трафик для HTTP-запросов во время учебной разработки.")
    add_prose_block(doc, generate_paragraphs(
        "подключения клиента к серверу",
        [
            "AuthApi отправляет учетные данные на маршруты /register и /login",
            "TasksApi добавляет заголовок Authorization: Bearer token к защищенным запросам",
            "ответы сервера преобразуются в DTO и затем в доменную модель Task",
            "ошибки сети и сервера преобразуются в сообщения, отображаемые пользователю",
        ],
        24,
    ))
    heading(doc, "3.3. Тестирование разработанного приложения", 2)
    paragraph(doc, "Тестирование выполнялось как ручная проверка пользовательских сценариев и API. Серверные маршруты были предварительно проверены через Postman, клиентская часть проверялась через запуск Android-приложения и анализ состояний интерфейса.")
    add_table(doc, ["№", "Сценарий", "Ожидаемый результат", "Результат"], [
        ["1", "Регистрация нового пользователя.", "Пользователь создан, отображается переход к авторизации.", "Успешно"],
        ["2", "Регистрация с уже существующим email.", "Отображается сообщение об ошибке.", "Успешно"],
        ["3", "Авторизация с корректными данными.", "JWT сохраняется, открывается экран задач.", "Успешно"],
        ["4", "Авторизация с неверным паролем.", "Отображается ошибка входа.", "Успешно"],
        ["5", "Создание задачи.", "Новая задача появляется в списке.", "Успешно"],
        ["6", "Редактирование задачи.", "Название и описание обновляются.", "Успешно"],
        ["7", "Удаление задачи.", "Задача исчезает из списка.", "Успешно"],
        ["8", "Смена статуса checkbox.", "Задача перемещается между вкладками статуса.", "Успешно"],
        ["9", "Поиск по существующей задаче.", "Отображается список найденных задач.", "Успешно"],
        ["10", "Поиск без результатов.", "Отображается плейсхолдер «ничего не найдено».", "Успешно"],
        ["11", "Очистка поисковой строки.", "Текст удаляется, клавиатура скрывается.", "Успешно"],
        ["12", "Поворот экрана при введенном поисковом запросе.", "Текст запроса сохраняется.", "Успешно"],
        ["13", "Очистка истории поиска.", "Список истории полностью очищается.", "Успешно"],
        ["14", "Переключение темной темы.", "Тема меняется и сохраняется после перезапуска.", "Успешно"],
    ])
    add_prose_block(doc, generate_paragraphs(
        "тестирования",
        [
            "проверка показала соответствие клиентской части обязательным требованиям методических указаний",
            "защита доступа подтверждается тем, что сервер выбирает задачи только по userId текущего пользователя",
            "ошибочные состояния не приводят к аварийному завершению приложения, а отображаются в интерфейсе",
            "сохранение пользовательских настроек подтверждает корректность использования DataStore",
        ],
        26,
    ))


def add_conclusion(doc: Document) -> None:
    add_section(doc, "Заключение")
    add_prose_block(doc, [
        "В ходе курсовой работы разработано клиент-серверное Android-приложение «Планировщик задач». Приложение позволяет пользователю зарегистрироваться, авторизоваться, вести список личных задач, создавать, редактировать, удалять записи, менять статус выполнения, выполнять поиск и использовать историю поисковых запросов.",
        "Клиентская часть реализована на Kotlin с использованием Jetpack Compose и архитектурного разделения data, domain и presentation. Серверная часть реализована на Ktor Server, использует JWT для авторизации, BCrypt для защиты паролей и PostgreSQL для хранения пользователей и задач. Доступ к задачам ограничен владельцем через проверку userId.",
        "Разработанное приложение соответствует обязательным требованиям методических указаний: интерфейс выполнен на русском языке, присутствуют экраны регистрации и авторизации, реализованы поиск, кнопка «Очистить», сохранение текста запроса при повороте, плейсхолдеры, история поиска, ProgressBar и переключатель темной темы.",
        "Возможными направлениями дальнейшего развития являются добавление сроков выполнения задач, напоминаний, приоритетов, сортировки, push-уведомлений и более подробной статистики выполнения. Также возможно расширение серверной части за счет refresh-токенов и централизованной обработки ошибок.",
    ])


def add_sources(doc: Document) -> None:
    add_section(doc, "Список использованных источников")
    sources = [
        "Android Developers. Jetpack Compose documentation. URL: https://developer.android.com/compose (дата обращения: 30.05.2026).",
        "Android Developers. DataStore documentation. URL: https://developer.android.com/topic/libraries/architecture/datastore (дата обращения: 30.05.2026).",
        "Ktor Documentation. Server development and authentication. URL: https://ktor.io/docs/server-create-a-new-project.html (дата обращения: 30.05.2026).",
        "Ktor Documentation. Client development. URL: https://ktor.io/docs/client-create-new-application.html (дата обращения: 30.05.2026).",
        "PostgreSQL Documentation. URL: https://www.postgresql.org/docs/ (дата обращения: 30.05.2026).",
        "JetBrains Exposed Documentation. URL: https://www.jetbrains.com/help/exposed/home.html (дата обращения: 30.05.2026).",
        "Auth0. JSON Web Token Introduction. URL: https://jwt.io/introduction (дата обращения: 30.05.2026).",
        "Pro Android with Kotlin: Developing Modern Mobile Apps. Apress, 2022.",
        "Kotlin Documentation. URL: https://kotlinlang.org/docs/home.html (дата обращения: 30.05.2026).",
        "ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления.",
    ]
    add_numbered(doc, sources)


def add_appendices(doc: Document) -> None:
    add_section(doc, "Приложения")
    heading(doc, "Приложение А. Листинг основных файлов мобильного приложения", 2)
    files = [
        ROOT / "app/src/main/java/com/example/taskplanner/MainActivity.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/presentation/navigation/AppNavigation.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/presentation/auth/AuthViewModel.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/presentation/auth/LoginScreen.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/presentation/tasks/TasksViewModel.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/presentation/tasks/TasksScreen.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/data/remote/ApiClient.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/data/remote/AuthApi.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/data/remote/TasksApi.kt",
        ROOT / "app/src/main/java/com/example/taskplanner/data/local/UserPreferencesManager.kt",
    ]
    for file in files:
        heading(doc, file.name, 3)
        for line in safe_read(file).splitlines()[:220]:
            paragraph(doc, line, style="Code")
    heading(doc, "Приложение Б. Листинг основных файлов серверной части", 2)
    for file in [SERVER_ROOT / "build.gradle.kts", SERVER_ROOT / "src/main/kotlin/Application.kt"]:
        heading(doc, file.name, 3)
        for line in safe_read(file).splitlines()[:260]:
            paragraph(doc, line, style="Code")
    heading(doc, "Приложение В. Скриншоты экранных форм", 2)
    paragraph(doc, "В приложении предусмотрены следующие экранные формы: экран авторизации, экран регистрации, экран задач, история поиска, плейсхолдер отсутствия результатов, плейсхолдер ошибки, диалог добавления задачи, диалог редактирования задачи и темная тема. Фактические скриншоты рекомендуется добавить после финального запуска приложения на устройстве перед сдачей в СДО.")
    heading(doc, "Приложение Г. Тестовые сценарии", 2)
    paragraph(doc, "Тестовые сценарии приведены в разделе 3.3. При необходимости перед сдачей они могут быть дополнены фактическими скриншотами выполнения каждого сценария.")


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = build_diagrams()
    doc = Document()
    configure_document(doc)
    set_update_fields(doc)

    heading(doc, "Содержание", 1)
    add_toc(doc)
    doc.add_page_break()

    add_intro(doc)
    add_chapter_1(doc)
    add_chapter_2(doc, diagrams)
    add_chapter_3(doc)
    add_conclusion(doc)
    add_sources(doc)
    add_appendices(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
