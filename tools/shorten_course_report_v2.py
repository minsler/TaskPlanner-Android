from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path(r"C:\Users\marya\Documents\х\курсач\Курсовая работа_TaskPlanner_без титульного листа.docx")
OUT = Path(r"C:\Users\marya\Documents\х\курсач\Курсовая работа_TaskPlanner_сокращенная_v2.docx")


SECTION_LIMITS = {
    "ВВЕДЕНИЕ": 5,
    "1.1.": 14,
    "1.2.": 10,
    "1.3.": 12,
    "2.1.": 12,
    "2.2.": 12,
    "2.3.1.": 1,
    "2.3.2.": 10,
    "3.1.": 16,
    "3.2.1.": 12,
    "3.2.2.": 10,
    "3.2.3.": 12,
    "3.3.": 12,
    "ЗАКЛЮЧЕНИЕ": 4,
}

REFERENCE_TEXTS = {
    "Таблица 1.3": "Выбор инструментальных средств для разработки приложения представлен в таблице 1.3.",
    "Таблица 2.1.1": "Нефункциональные требования к мобильному приложению представлены в таблице 2.1.1.",
    "Таблица 2.1": "Функциональные требования к мобильному приложению представлены в таблице 2.1.",
    "Таблица 2.3.2": "Структура таблиц базы данных представлена в таблице 2.3.2.",
    "Таблица 3.1": "Структура клиентской части приложения представлена в таблице 3.1.",
    "Таблица 3.2.1": "Перечень серверных маршрутов и их назначение представлены в таблице 3.2.1.",
    "Таблица 3.3": "Результаты проверки пользовательских сценариев представлены в таблице 3.3.",
    "Рисунок 2.1": "Use case диаграмма приложения представлена на рисунке 2.1.",
    "Рисунок 2.3.1": "Диаграмма классов клиентской части представлена на рисунке 2.3.1.",
    "Рисунок 2.5": "IDEF0-диаграмма процесса управления задачами представлена на рисунке 2.5.",
    "Рисунок 2.3.2": "Логическая модель базы данных представлена на рисунке 2.3.2.",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def is_appendix_heading(text: str, style: str) -> bool:
    return style.startswith("Heading 1") and text.strip().upper() == "ПРИЛОЖЕНИЯ"


def section_key(text: str) -> str | None:
    stripped = text.strip()
    upper = stripped.upper()
    if upper in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
        return upper
    for prefix in (
        "1.1.", "1.2.", "1.3.",
        "2.1.", "2.2.", "2.3.1.", "2.3.2.",
        "3.1.", "3.2.1.", "3.2.2.", "3.2.3.", "3.3.",
    ):
        if stripped.startswith(prefix):
            return prefix
    return None


def is_plain_prose(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph.style.name != "Normal":
        return False
    if text.startswith("Таблица") or text.startswith("Рисунок"):
        return False
    return True


def remove_screenshot_figure(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Интерфейс приложения, представленный на рисунке 2.2.1"):
            paragraph.text = (
                "Интерфейс приложения построен на нескольких основных экранных формах: экран авторизации, "
                "экран регистрации, экран списка задач, диалоги добавления и редактирования задачи. "
                "Межэкранное взаимодействие реализовано средствами Navigation Compose и представлено на рисунке 2.2."
            )
        if text.startswith("Рисунок 2.2.1"):
            previous = paragraph._element.getprevious()
            if previous is not None and previous.tag == qn("w:p"):
                paragraph._element.getparent().remove(previous)
            delete_paragraph(paragraph)
            break


def shorten_before_appendices(doc: Document) -> None:
    current = None
    counts: dict[str, int] = {}
    to_delete = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name
        if is_appendix_heading(text, style):
            break

        if style.startswith("Heading"):
            key = section_key(text)
            if key is not None:
                current = key
                counts.setdefault(current, 0)
            continue

        if current and is_plain_prose(paragraph):
            counts[current] = counts.get(current, 0) + 1
            if counts[current] > SECTION_LIMITS.get(current, 999):
                to_delete.append(paragraph)

    for paragraph in to_delete:
        delete_paragraph(paragraph)


def previous_paragraph_texts(paragraph, limit: int = 4) -> str:
    element = paragraph._element.getprevious()
    texts = []
    while element is not None and len(texts) < limit:
        if element.tag == qn("w:p"):
            texts.append("".join(node.text or "" for node in element.iter() if node.tag == qn("w:t")))
        element = element.getprevious()
    return "\n".join(texts)


def reference_for_caption(text: str) -> str | None:
    for marker, reference in REFERENCE_TEXTS.items():
        if text.startswith(marker):
            return reference
    return None


def add_reference_before(paragraph, reference: str, before_image: bool = False) -> None:
    new_p = paragraph.insert_paragraph_before(reference)
    new_p.style = "Normal"
    if before_image:
        image_element = paragraph._element.getprevious()
        if image_element is not None and image_element.tag == qn("w:p"):
            image_element.addprevious(new_p._element)


def has_recent_reference(paragraph, caption_text: str) -> bool:
    match = re.search(r"((?:Таблица|Рисунок)\s+\d+(?:\.\d+)*)", caption_text)
    if not match:
        return False
    number = match.group(1)
    previous = previous_paragraph_texts(paragraph)
    return number in previous and "представ" in previous.lower()


def add_missing_references(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if is_appendix_heading(text, paragraph.style.name):
            break

        reference = reference_for_caption(text)
        if not reference or has_recent_reference(paragraph, text):
            continue

        if text.startswith("Рисунок"):
            add_reference_before(paragraph, reference, before_image=True)
        else:
            add_reference_before(paragraph, reference)


def normalize_toc_styles(doc: Document) -> None:
    for name in ("TOC 1", "TOC 2", "TOC 3", "toc 1", "toc 2", "toc 3"):
        if name not in doc.styles:
            continue
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    remove_screenshot_figure(doc)
    shorten_before_appendices(doc)
    add_missing_references(doc)
    normalize_toc_styles(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
