from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


SRC = Path(r"C:\Users\marya\Documents\х\курсач\Курсовая работа_TaskPlanner_без титульного листа.docx")
OUT = Path(r"C:\Users\marya\Documents\х\курсач\Курсовая работа_TaskPlanner_сокращенная.docx")


SECTION_LIMITS = {
    "ВВЕДЕНИЕ": 5,
    "1.1.": 5,
    "1.2.": 4,
    "1.3.": 5,
    "2.1.": 5,
    "2.2.": 5,
    "2.3.1.": 5,
    "2.3.2.": 4,
    "3.1.": 6,
    "3.2.1.": 5,
    "3.2.2.": 4,
    "3.2.3.": 4,
    "3.3.": 5,
    "ЗАКЛЮЧЕНИЕ": 4,
}

REFERENCE_TEXTS = {
    "Таблица 1.2": "Сравнительный анализ существующих аналогов представлен в таблице 1.2.",
    "Таблица 1.3": "Обоснование выбора инструментальных средств представлено в таблице 1.3.",
    "Таблица 2.1 ": "Функциональные требования к приложению представлены в таблице 2.1.",
    "Таблица 2.1.1": "Нефункциональные требования к приложению представлены в таблице 2.1.1.",
    "Таблица 2.3.2": "Описание таблиц базы данных представлено в таблице 2.3.2.",
    "Таблица 3.1": "Описание компонентов клиентской части представлено в таблице 3.1.",
    "Таблица 3.2.1": "Основная логика серверных маршрутов представлена в таблице 3.2.1.",
    "Таблица 3.3": "Тестовые сценарии приложения представлены в таблице 3.3.",
    "Рисунок 2.1": "Use case диаграмма приложения представлена на рисунке 2.1.",
    "Рисунок 2.2 ": "Схема межэкранного взаимодействия представлена на рисунке 2.2.",
    "Рисунок 2.2.1": "Ключевые экраны мобильного приложения представлены на рисунке 2.2.1.",
    "Рисунок 2.3.1": "Диаграмма классов клиентской части представлена на рисунке 2.3.1.",
    "Рисунок 2.5": "Диаграмма IDEF0 представлена на рисунке 2.5.",
    "Рисунок 2.3.2": "Логическая модель базы данных представлена на рисунке 2.3.2.",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def is_appendix_heading(text: str, style: str) -> bool:
    return style.startswith("Heading 1") and text.strip().upper() == "ПРИЛОЖЕНИЯ"


def section_key(text: str) -> str | None:
    upper = text.strip().upper()
    if upper in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
        return upper
    for prefix in (
        "1.1.", "1.2.", "1.3.",
        "2.1.", "2.2.", "2.3.1.", "2.3.2.",
        "3.1.", "3.2.1.", "3.2.2.", "3.2.3.", "3.3.",
    ):
        if text.strip().startswith(prefix):
            return prefix
    return None


def is_plain_prose(paragraph) -> bool:
    text = paragraph.text.strip()
    if not text:
        return False
    if paragraph.style.name != "Normal":
        return False
    if text.startswith("Таблица") or text.startswith("Рисунок"):
        return False
    if "представлен" in text.lower() and ("таблиц" in text.lower() or "рисунк" in text.lower()):
        return False
    return True


def shorten_before_appendices(doc: Document) -> None:
    current = None
    counts: dict[str, int] = {}
    to_delete = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style = paragraph.style.name
        if is_appendix_heading(text, style):
            break

        if style.startswith("Heading"):
            key = section_key(text)
            if key is not None:
                current = key
                counts.setdefault(current, 0)
            continue

        if current and is_plain_prose(paragraph):
            counts[current] = counts.get(current, 0) + 1
            if counts[current] > SECTION_LIMITS.get(current, 999):
                to_delete.append(paragraph)

    for paragraph in to_delete:
        delete_paragraph(paragraph)


def previous_text(paragraph) -> str:
    element = paragraph._element.getprevious()
    texts = []
    while element is not None and len(texts) < 3:
        if element.tag == qn("w:p"):
            texts.append("".join(node.text or "" for node in element.iter() if node.tag == qn("w:t")))
        element = element.getprevious()
    return "\n".join(texts)


def find_reference(text: str) -> str | None:
    for marker, reference in REFERENCE_TEXTS.items():
        if text.startswith(marker):
            return reference
    return None


def add_reference_before(paragraph, text: str) -> None:
    new_p = paragraph.insert_paragraph_before(text)
    new_p.style = "Normal"


def add_references(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if is_appendix_heading(text, paragraph.style.name):
            break

        reference = find_reference(text)
        if not reference:
            continue

        prev = previous_text(paragraph)
        number_match = re.search(r"((?:Таблица|Рисунок)\s+\d+(?:\.\d+)*)", text)
        number = number_match.group(1) if number_match else ""
        if number and number in prev and ("представлен" in prev.lower() or "представлена" in prev.lower()):
            continue

        if text.startswith("Рисунок"):
            image_element = paragraph._element.getprevious()
            if image_element is not None and image_element.tag == qn("w:p"):
                # Insert the reference before the image paragraph, not between the image and caption.
                new_p = paragraph.insert_paragraph_before(reference)
                new_p.style = "Normal"
                image_element.addprevious(new_p._element)
            else:
                add_reference_before(paragraph, reference)
        else:
            add_reference_before(paragraph, reference)


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(OUT)
    shorten_before_appendices(doc)
    add_references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
